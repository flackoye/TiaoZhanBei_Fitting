from __future__ import annotations

import hashlib
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SOURCE = Path(r"E:\DeskTop\preparation\Chen_Wei\TiaoZhanBei\5.5.docx")
OUTPUT = Path(
    r"E:\DeskTop\preparation\Chen_Wei\TiaoZhanBei\Fitting\outputs\documents\5.5_反演结果不确定性与可信度评价_修订版.docx"
)

BODY_CN = "宋体"
HEADING_CN = "黑体"
LATIN = "Times New Roman"
MATH = "Cambria Math"
ACCENT = "1F4E78"
HEADER_FILL = "D9EAF7"
ALT_FILL = "F5F9FC"


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for block in iter(lambda: f.read(1024 * 1024), b""):
            h.update(block)
    return h.hexdigest()


def set_run_font(run, cn=BODY_CN, latin=LATIN, size=Pt(11), bold=None, color=None):
    run.font.name = latin
    run.font.size = size
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), cn)
    rfonts.set(qn("w:cs"), latin)


def clear_body(doc: Document):
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def configure_styles(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = LATIN
    normal.font.size = Pt(11)
    normal._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), BODY_CN)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.first_line_indent = Pt(22)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)

    for name, size, before, after in [
        ("Heading 2", 16, 13, 10),
        ("Heading 3", 16, 12, 8),
        ("Heading 4", 14, 10, 6),
    ]:
        style = doc.styles[name]
        style.font.name = LATIN
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        rpr = style._element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        rfonts.set(qn("w:ascii"), LATIN)
        rfonts.set(qn("w:hAnsi"), LATIN)
        rfonts.set(qn("w:eastAsia"), HEADING_CN)
        style.paragraph_format.first_line_indent = Pt(0)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True


def add_body(doc: Document, text: str, *, indent=True, after=Pt(6)):
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Pt(22) if indent else Pt(0)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = after
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_heading(doc: Document, text: str, level: int):
    style = "Heading 2" if level == 2 else "Heading 3" if level == 3 else "Heading 4"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.first_line_indent = Pt(0)
    r = p.add_run(text)
    set_run_font(
        r,
        cn=HEADING_CN,
        latin=LATIN,
        size=Pt(16 if level in (2, 3) else 14),
        bold=True,
    )
    return p


def add_equation(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_together = True
    r = p.add_run(text)
    set_run_font(r, cn=MATH, latin=MATH, size=Pt(11))
    return p


def add_label_paragraph(doc: Document, label: str, text: str):
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.left_indent = Pt(22)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(label)
    set_run_font(r1, bold=True)
    r2 = p.add_run(text)
    set_run_font(r2)
    return p


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "6")
        tag.set(qn("w:color"), "8EA9BF")


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_table_width(table, width_twips):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_twips))
    tbl_w.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def set_table_grid(table, widths):
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)


def set_cell_width(cell, width_twips):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def add_table(doc: Document, headers, rows, widths, *, numeric_cols=()):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Normal Table"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    total = sum(widths)
    set_table_width(table, total)
    set_table_grid(table, widths)
    set_table_borders(table)
    hdr = table.rows[0]
    set_repeat_header(hdr)

    for j, (cell, text, width) in enumerate(zip(hdr.cells, headers, widths)):
        set_cell_width(cell, width)
        set_cell_margins(cell)
        set_cell_shading(cell, HEADER_FILL)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(str(text))
        set_run_font(r, cn=HEADING_CN, size=Pt(9.5), bold=True)

    for i, row in enumerate(rows):
        cells = table.add_row().cells
        for j, (cell, text, width) in enumerate(zip(cells, row, widths)):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            if i % 2 == 1:
                set_cell_shading(cell, ALT_FILL)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j in numeric_cols else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            r = p.add_run(str(text))
            set_run_font(r, size=Pt(9.2))

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    spacer.paragraph_format.line_spacing = 1.0
    return table


def build_document():
    original_hash = sha256(SOURCE)
    doc = Document(SOURCE)
    clear_body(doc)
    configure_styles(doc)

    add_heading(doc, "5.5 反演结果不确定性与可信度评价", 2)
    add_body(
        doc,
        "在深部强动压巷道围岩随钻感知与智能卸压调控体系中，前述反演模型能够沿钻孔轨迹连续输出围岩损伤程度和应力水平。然而，井下机械噪声、排屑脉动、钻头磨损、工况变化以及模型训练样本覆盖范围均会使点预测产生不同程度的误差。若后续曲线拟合和空间重构对所有测点赋予相同信任程度，局部异常或高风险预测可能在插值过程中被放大，进而影响危险区域定位。",
    )
    add_body(
        doc,
        "为此，本节以三个反演模型的逐点预测、模型状态置信度及真实标签为基础，构建面向单孔反演结果的不确定性评价方法。研究同时设计条件分位残差回归、模型置信度单调分位校准和多证据稳健融合三套候选方案，通过留一数据源交叉验证比较其误差风险排序能力，并将最优方案输出的损伤权重和应力权重输入单孔曲线拟合。最终形成“多模型预测聚合—不确定性估计—样本外评价—可信度赋权—曲线拟合”的完整处理链，为 5.6 节二维、三维围岩状态场重构提供带有测点级可信度约束的基础数据。",
    )

    add_heading(doc, "5.5.1 不确定性来源分析", 3)
    add_body(
        doc,
        "单孔随钻反演不确定性主要由数据不确定性和模型不确定性共同构成。数据不确定性来源于传感器噪声、电信号干扰、排屑不畅、卡钻、钻头磨损及局部地质界面等因素，通常表现为原始响应的随机波动、脉冲毛刺、短时漂移或数据缺失。这类扰动会沿特征提取和状态反演过程传递，使相邻孔深位置的预测结果出现非物理跳变。",
    )
    add_body(
        doc,
        "模型不确定性来源于反演模型的结构差异、参数估计误差和训练样本边界。当当前工况与训练样本相似时，不同模型通常给出相近预测；当工况复杂或信息不足时，模型之间可能出现分歧，模型内部状态置信度也可能下降。但模型高度同质化时，即使预测发生偏差，不同模型仍可能同时给出相近结果，因此模型间分歧不能单独作为可靠的不确定性指标。",
    )
    add_body(
        doc,
        "本节只评价单孔沿孔深反演结果的测点级可信程度，并分别形成损伤权重和应力权重。钻孔数量、周向覆盖、轨迹定位误差、空间插值距离及无观测区域造成的空间重构不确定性不在本节提前合并，而由 5.6 节在单孔状态值和单孔权重基础上继续计算，从而保持单孔反演可信度与空间重构可信度的职责边界。",
    )

    add_heading(doc, "5.5.2 多模型预测聚合与三套候选算法", 3)
    add_body(
        doc,
        "实验数据来源于 all_model_predictions.csv，共包含 286,243 行模型预测记录。按照 source_file、experiment_id 和 sample_index 对同一空间位置的三个模型预测进行聚合后，得到 96,281 个位置级测点，覆盖 3 个反演模型和 13 个独立数据源。损伤和应力分别进行聚合与不确定性建模。",
    )
    add_table(
        doc,
        ["数据项目", "数量或设置"],
        [
            ["模型预测记录", "286,243 行"],
            ["位置级测点", "96,281 个"],
            ["反演模型", "3 个"],
            ["独立数据源", "13 个 source_file"],
            ["位置聚合键", "source_file / experiment_id / sample_index"],
            ["评价目标", "损伤程度、应力水平分别处理"],
        ],
        [3900, 4400],
    )
    add_body(
        doc,
        "设第 i 个测点上第 m 个模型对目标 k 的预测为 ŷᵢₘ⁽ᵏ⁾，其中 k 分别表示损伤或应力。为降低单个模型极端值对参考预测的影响，位置级集成预测采用三个模型预测的中位数，同时计算模型间标准差、预测极差和平均状态置信度。",
    )
    add_equation(doc, "ỹᵢ⁽ᵏ⁾ = medianₘ{ŷᵢₘ⁽ᵏ⁾}")
    add_equation(doc, "σᵢ⁽ᵏ⁾ = Stdₘ{ŷᵢₘ⁽ᵏ⁾}，    Rᵢ⁽ᵏ⁾ = maxₘŷᵢₘ⁽ᵏ⁾ − minₘŷᵢₘ⁽ᵏ⁾")
    add_equation(doc, "c̄ᵢ = (1/M)Σₘ₌₁ᴹ cᵢₘ，    eᵢ⁽ᵏ⁾ = |ỹᵢ⁽ᵏ⁾ − yᵢ⁽ᵏ⁾|")
    add_body(
        doc,
        "其中，ỹᵢ⁽ᵏ⁾ 为位置级集成中位数，σᵢ⁽ᵏ⁾ 为模型间标准差，Rᵢ⁽ᵏ⁾ 为预测极差。实验阶段以集成中位数与真实标签之差的绝对值 eᵢ⁽ᵏ⁾=|ỹᵢ⁽ᵏ⁾−yᵢ⁽ᵏ⁾| 作为不确定性算法的评价对象。",
    )

    add_heading(doc, "（1）方案一：条件分位残差回归", 4)
    add_body(
        doc,
        "方案一将模型间标准差、预测极差、平均状态置信度、累计孔深和区段编号组成五维特征向量，采用 HistGradientBoostingRegressor 建立 90% 分位残差回归模型，直接预测给定特征条件下绝对误差的高分位风险。其目标不是重新预测损伤或应力，而是对现有反演结果可能产生的误差进行监督校准。",
    )
    add_equation(doc, "xᵢ⁽ᵏ⁾ = [σᵢ⁽ᵏ⁾, Rᵢ⁽ᵏ⁾, c̄ᵢ, dᵢ, sᵢ]")
    add_equation(doc, "ρτ(r) = τr  (r ≥ 0)；    ρτ(r) = (τ−1)r  (r < 0)")
    add_equation(doc, "q̂₀.₉ = arg min_q Σᵢ ρ₀.₉(eᵢ⁽ᵏ⁾ − q(xᵢ⁽ᵏ⁾))")
    add_equation(doc, "uᵢ,₁⁽ᵏ⁾ = max{0, q̂₀.₉(eᵢ⁽ᵏ⁾ | xᵢ⁽ᵏ⁾)}")
    add_body(
        doc,
        "其中，c̄ᵢ 为三个模型状态置信度的平均值，dᵢ 为累计孔深，sᵢ 为区段编号，uᵢ,₁⁽ᵏ⁾ 为条件高分位误差风险。该值反映相似历史样本中的较大误差水平，不解释为预测正确概率。",
    )

    add_heading(doc, "（2）方案二：模型置信度单调分位校准", 4)
    add_body(
        doc,
        "方案二以模型平均状态置信度为单一核心信号，定义置信缺失程度 zᵢ=1−c̄ᵢ。训练数据按 zᵢ 的分位数划分为 15 个区间，在每个区间内统计绝对误差的 90% 分位数，再采用保序回归建立非递减校准函数，使置信缺失程度增加时估计风险原则上不下降。",
    )
    add_equation(doc, "zᵢ = 1 − c̄ᵢ，    uᵢ,₂⁽ᵏ⁾ = gₖ(zᵢ)，    gₖ(zₐ) ≤ gₖ(zᵦ)  (zₐ < zᵦ)")
    add_equation(doc, "qᵦ⁽ᵏ⁾ = Q₀.₉{eᵢ⁽ᵏ⁾ | zᵢ ∈ Bᵦ}")
    add_equation(doc, "gₖ = arg min_g Σᵦ nᵦ[qᵦ⁽ᵏ⁾ − g(z̄ᵦ)]²，    且 g 单调非递减")
    add_body(
        doc,
        "该方案结构简单、单调关系清晰，能够检验模型自身置信度是否具有独立的误差指示能力；但它难以识别模型高置信度错误，也不能直接利用孔深、区段及模型分歧等补充信息。",
    )

    add_heading(doc, "（3）方案三：多证据稳健融合", 4)
    add_body(
        doc,
        "方案三分别构造模型分歧、置信缺失和局部不稳定三个百分位分量。局部不稳定性采用滑动窗口中位数和中位绝对偏差评价，并对连续不少于三个测点、多个模型表现一致的持续变化实施边界保护，避免将真实地质界面简单判为低可信异常。三个分量按非负系数加权融合，系数以 0.2 为步长搜索，共形成 21 组满足系数和为 1 的候选组合。",
    )
    add_equation(doc, "sᵢ⁽ᵏ⁾ = αₖaᵢ⁽ᵏ⁾ + βₖbᵢ + γₖc̃ᵢ⁽ᵏ⁾")
    add_equation(doc, "αₖ + βₖ + γₖ = 1，    αₖ, βₖ, γₖ ≥ 0")
    add_equation(doc, "aᵢ⁽ᵏ⁾ = 0.5[Pσ(σᵢ⁽ᵏ⁾) + PR(Rᵢ⁽ᵏ⁾)]，    bᵢ = Pc(1−c̄ᵢ)")
    add_equation(doc, "mᵢ⁽ᵏ⁾ = medianⱼ∈Nᵢ{ỹⱼ⁽ᵏ⁾}，    MADᵢ⁽ᵏ⁾ = medianⱼ∈Nᵢ|ỹⱼ⁽ᵏ⁾−mᵢ⁽ᵏ⁾|")
    add_equation(doc, "rᵢ⁽ᵏ⁾ = |ỹᵢ⁽ᵏ⁾−mᵢ⁽ᵏ⁾| / [1.4826MADᵢ⁽ᵏ⁾+ε]")
    add_equation(doc, "c̃ᵢ⁽ᵏ⁾ = hᵢ⁽ᵏ⁾Pr(rᵢ⁽ᵏ⁾)，    hᵢ⁽ᵏ⁾ = η（受保护边界）或 1（其他位置）")
    add_body(
        doc,
        "其中，aᵢ⁽ᵏ⁾ 表示模型分歧百分位，bᵢ 表示置信缺失百分位，c̃ᵢ⁽ᵏ⁾ 表示经过真实边界保护的局部不稳定百分位。该方案旨在通过多源证据互补识别单一指标遗漏的风险位置。",
    )

    add_heading(doc, "5.5.3 留一数据源交叉验证与评价方法", 3)
    add_body(
        doc,
        "为避免同一数据源中的相邻测点同时进入训练集和测试集造成信息泄漏，采用 Leave-one-source_file-out 方式开展 13 折交叉验证。每一折保留 1 个 source_file 作为测试数据，其余 12 个数据源用于训练或确定校准关系。损伤和应力分别训练、分别评价，三套方案使用相同的数据划分和绝对误差目标。",
    )
    add_body(
        doc,
        "不确定性算法首先按照估计风险由低到高排列测试测点。随着覆盖率逐步增加，计算当前保留测点的累计平均绝对误差，形成风险—覆盖曲线，其曲线下面积定义为 AURC。AURC 越低，说明低不确定性样本的实际误差越小，风险排序能力越强。",
    )
    add_equation(doc, "Risk(t) = (1 / nₜ) Σᵢ₌₁ⁿᵗ e₍ᵢ₎，    AURC = ∫₀¹ Risk(c) dc")
    add_equation(doc, "ρₛ = 1 − [6Σᵢ₌₁ᴺ dᵢ²] / [N(N²−1)]")
    add_body(
        doc,
        "同时计算不确定性与绝对误差的 Spearman 秩相关系数。Spearman 越大，表示不确定性随实际误差增大而上升的趋势越明显。方法选择以 13 折平均 AURC 为主要依据，以平均 Spearman 为辅助依据。方案三的 21 组系数分别在相同 13 折上评价，先选出跨折平均表现最好的固定系数组合，再与方案一、方案二比较；因此不能把 273 条“系数组合×交叉验证折”记录直接视为 273 个独立折。",
    )

    add_heading(doc, "5.5.4 三套方案实验结果与最优方法确定", 3)
    add_body(
        doc,
        "三套方案的样本外交叉验证结果如表所示。损伤和应力目标上，条件分位残差回归均取得最低 AURC 和最高 Spearman，说明该方法能够更有效地把实际误差较小的测点排在低风险区域，并对误差变化保持更稳定的正向响应。",
    )
    add_table(
        doc,
        ["目标", "不确定性方案", "平均AURC↓", "平均Spearman↑", "融合系数(α,β,γ)", "结论"],
        [
            ["损伤", "方案一 条件分位残差回归", "7.39", "0.382", "—", "选取"],
            ["损伤", "方案二 单调分位校准", "10.73", "0.116", "—", "未选"],
            ["损伤", "方案三 多证据融合", "9.02", "0.186", "(0.8, 0.2, 0.0)", "未选"],
            ["应力", "方案一 条件分位残差回归", "1.76", "0.378", "—", "选取"],
            ["应力", "方案二 单调分位校准", "3.09", "0.095", "—", "未选"],
            ["应力", "方案三 多证据融合", "3.08", "0.040", "(0.0, 1.0, 0.0)", "未选"],
        ],
        [700, 2250, 1000, 1150, 1750, 700],
        numeric_cols=(0, 2, 3, 4, 5),
    )
    add_body(
        doc,
        "方案二未被选取的主要原因是状态置信度信号本身存在明显偏斜。其均值约为 0.83、中位数约为 0.92，大量测点集中于高置信度区间，且 33.2% 的状态置信度缺失值全部来自 advancedV1 模型。虽然缺失值按中性值处理能够维持流程运行，但仅依赖单一置信度信号难以区分高置信区域内部的细微误差差异，因此两个目标的 AURC 和 Spearman 均弱于方案一。",
    )
    add_body(
        doc,
        "方案三未被选取则反映了多证据并不必然带来有效增益。损伤和应力分别有 54.2% 和 76.2% 的位置模型间标准差为零，说明三个模型在多数位置给出完全一致的预测，模型分歧分量的信息量有限。进一步地，损伤最佳融合系数为 (0.8, 0.2, 0.0)，应力最佳组合为 (0.0, 1.0, 0.0)，两个目标均令局部不稳定分量 γ 为零，表明在当前数据条件下，MAD 局部偏离度没有提供稳定的额外误差排序信息。",
    )
    add_body(
        doc,
        "方案一能够同时使用模型分歧、平均置信度、累计孔深和区段编号五项特征。即使标准差、极差或置信度中的部分信号退化，梯度提升树仍可从剩余特征及其非线性关系中学习误差风险。因此，最终损伤和应力均选择条件分位残差回归作为当前数据集上的测点级不确定性估计方法。",
    )

    add_heading(doc, "5.5.5 最优方案的可信度及权重构建", 3)
    add_body(
        doc,
        "为将误差风险转换为下游算法可直接使用的权重，依据训练参考分布计算不确定性的经验百分位 pᵢ⁽ᵏ⁾，并将其反向映射为相对可信度 Cᵢ⁽ᵏ⁾。不确定性越小，其百分位越低，相对可信度和拟合权重越高。最低权重设置为 0.1，避免单个测点在后续处理中被完全删除。",
    )
    add_equation(doc, "pᵢ⁽ᵏ⁾ = #{uⱼ⁽ᵏ⁾ ≤ uᵢ⁽ᵏ⁾} / N")
    add_equation(doc, "Cᵢ⁽ᵏ⁾ = 1 − pᵢ⁽ᵏ⁾，    wᵢ⁽ᵏ⁾ = 0.1 + 0.9Cᵢ⁽ᵏ⁾")
    add_body(
        doc,
        "最终分别输出 damage_uncertainty、damage_confidence、damage_weight 以及 stress_uncertainty、stress_confidence、stress_weight。两类目标独立赋权，避免用同一个综合权重掩盖损伤预测和应力预测在误差结构上的差异。方案一生成的权重分布如下。",
    )
    add_table(
        doc,
        ["目标权重", "均值", "标准差", "范围", "低权占比[0.10,0.30]", "高权占比[0.70,1.00]"],
        [
            ["damage_weight", "0.409", "0.229", "0.100～0.965", "23.0%", "15.0%"],
            ["stress_weight", "0.379", "0.203", "0.100～0.567", "44.8%", "0.0%"],
        ],
        [1550, 780, 850, 1300, 1800, 1800],
        numeric_cols=(1, 2, 3, 4, 5),
    )
    add_body(
        doc,
        "损伤权重覆盖 0.100～0.965，能够形成较明显的逐点差异化控制。应力权重最高仅为 0.567，且没有大于 0.7 的高权测点。这一现象不能解释为三个模型在应力预测上普遍不一致，因为 76.2% 的位置模型间标准差恰好为零。其直接原因是应力残差校准结果只有少量离散档位并存在大量并列值，经验分布映射后权重分辨率受到限制。因此，当前应力权重更适合解释为相对风险分级，而不是精细概率置信度。",
    )

    add_heading(doc, "5.5.6 可信度驱动的单孔曲线拟合", 3)
    add_body(
        doc,
        "获得损伤权重和应力权重后，分别按 model、source_file 和 experiment_id 对单孔序列分组，并按累计孔深排序。首先采用 Hampel 方法检测孤立异常点，对需要修正的异常值进行线性插值；连续不少于三个测点的异常段作为潜在真实边界保护，不进行简单抹平。随后采用 Savitzky-Golay 方法得到平滑参考曲线，并由逐测点权重控制修正后原值与平滑值之间的融合比例。",
    )
    add_equation(doc, "mᵢ = medianⱼ∈Nᵢ(yⱼ)，    MADᵢ = medianⱼ∈Nᵢ|yⱼ−mᵢ|")
    add_equation(doc, "Iᵢ = 1{|yᵢ−mᵢ| > 3.5×1.4826MADᵢ}")
    add_equation(doc, "λᵢ⁽ᵏ⁾ = 0.8 − 0.6wᵢ⁽ᵏ⁾")
    add_equation(doc, "ŷᵢ,fit⁽ᵏ⁾ = (1−λᵢ⁽ᵏ⁾)ŷᵢ,corrected⁽ᵏ⁾ + λᵢ⁽ᵏ⁾ŷᵢ,smooth⁽ᵏ⁾")
    add_body(
        doc,
        "高权重点对应较小的 λᵢ⁽ᵏ⁾，拟合结果更多保留修正后原始特征；低权重点对应较大的 λᵢ⁽ᵏ⁾，拟合结果更多向平滑曲线靠近。最后采用 PCHIP 在等间距孔深网格上进行保形致密插值，形成用于后续空间重构的连续损伤曲线和连续应力曲线。",
    )
    add_table(
        doc,
        ["目标", "原始TV", "有权重拟合TV", "无权重拟合TV", "有权重roughness", "无权重roughness", "过冲"],
        [
            ["损伤", "420", "442", "466", "0.554", "0.204", "无"],
            ["应力", "680", "419", "413", "0.191", "0.104", "无"],
        ],
        [750, 850, 1250, 1250, 1300, 1300, 650],
        numeric_cols=(0, 1, 2, 3, 4, 5, 6),
    )
    add_body(
        doc,
        "单组对比中，损伤采用不确定性权重后的拟合总变差为 442，低于无权重模式的 466，表明差异化权重在该组中能够减少过度平滑造成的特征改变。应力有权重和无权重模式的拟合总变差分别为 419 和 413，差异较小；两种模式均未产生范围过冲。由于当前直接对比只覆盖一个分组，该结果用于验证权重接口和自适应融合机制，并提示损伤特征保持存在积极趋势，尚不能据此宣称加权方案在全量数据上取得显著提升。",
    )

    add_heading(doc, "5.5.7 结果讨论与后续空间重构接口", 3)
    add_body(
        doc,
        "三套算法的统一实验表明，条件分位残差回归在损伤和应力目标上均获得最低平均 AURC 和最高平均 Spearman，是当前数据条件下更有效的测点级不确定性估计方案。单调分位校准验证了模型自身状态置信度具有一定风险指示作用，但单一信号受高值集中和缺失影响较大；多证据融合则揭示了模型分歧和局部不稳定性在当前数据中的有效信息有限。这两套未选方案并非无效尝试，而是通过对照实验明确了方案一获选的原因及现有数据的信号边界。",
    )
    add_body(
        doc,
        "当前方法仍存在三方面局限。第一，三个反演模型在大量位置给出完全一致的预测，模型多样性不足限制了集成分歧的表达能力；第二，应力校准不确定性只有少量离散档位，导致应力权重的动态范围和分辨率有限；第三，有权重与无权重的直接拟合比较目前只完成单组验证，后续仍需在全部模型—数据源—实验分组上开展配对统计，进一步评价真实误差、曲线连续性、峰值保持和边界保持之间的综合变化。",
    )
    add_body(
        doc,
        "本节最终向 5.6 节提供四类核心输入：沿累计孔深分布的连续损伤曲线、连续应力曲线、damage_weight 和 stress_weight。5.6 节在此基础上结合孔口位置、钻孔方位与倾角、实际轨迹、邻域钻孔覆盖、空间距离及插值不确定性，分别重构二维损伤状态场、二维应力状态场和三维空间可信度体。由此形成从单孔反演可信度到空间重构可信度的分层传递关系，避免将单孔模型风险与空间外推风险混为同一指标。",
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = "5.5 反演结果不确定性与可信度评价（修订版）"
    doc.core_properties.subject = "三套不确定性算法实验、方法选择及拟合权重接口"
    doc.core_properties.keywords = "不确定性, 条件分位残差回归, AURC, Spearman, 可信度权重"
    doc.save(OUTPUT)
    assert sha256(SOURCE) == original_hash, "Source DOCX was unexpectedly modified"
    print(f"Created: {OUTPUT}")
    print(f"Source SHA256 preserved: {original_hash}")
    print(f"Output bytes: {OUTPUT.stat().st_size}")


if __name__ == "__main__":
    build_document()
