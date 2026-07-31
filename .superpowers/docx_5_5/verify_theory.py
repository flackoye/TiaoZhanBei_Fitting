from pathlib import Path
from hashlib import sha256
from zipfile import ZipFile

from docx import Document
from lxml import etree


ROOT = Path(r"E:\DeskTop\preparation\Chen_Wei\TiaoZhanBei\Fitting")
SOURCE = ROOT / "outputs" / "documents" / "5.5_反演结果不确定性与可信度评价_修订版.docx"
OUTPUT = ROOT / "outputs" / "documents" / "5.5_反演结果不确定性与可信度评价_理论扩充版.docx"


def table_hash(table):
    return sha256(etree.tostring(table._tbl, encoding="utf-8")).hexdigest()


source = Document(SOURCE)
output = Document(OUTPUT)
source_hashes = [table_hash(t) for t in source.tables[:3]]
output_hashes = [table_hash(t) for t in output.tables]
assert len(source.tables) == 4, len(source.tables)
assert len(output.tables) == 3, len(output.tables)
assert output_hashes == source_hashes, (source_hashes, output_hashes)

paragraph_text = "\n".join(p.text for p in output.paragraphs)
table_text = "\n".join(
    cell.text
    for table in output.tables
    for row in table.rows
    for cell in row.cells
)
text = paragraph_text + "\n" + table_text
headings = [f"5.5.{i}" for i in range(1, 8)]
assert all(h in text for h in headings)

required = [
    "Var(y|x)", "TV(y)", "Roughness(y)", "ρτ", "AURC", "ρₛ",
    "286243", "96281", "7.39", "0.382", "1.76", "0.378",
    "0.409", "0.965",
    "0.379", "0.567", "420", "442", "466", "680", "419", "413",
    "0.554", "0.204", "0.191", "0.104",
]
missing = [item for item in required if item not in text.replace(",", "")]
assert not missing, missing
assert "(0.8, 0.2, 0.0)" in text
assert "(0.0, 1.0, 0.0)" in text

with ZipFile(OUTPUT) as archive:
    bad = archive.testzip()
    assert bad is None, bad

print(f"PASS: headings={len(headings)}, tables={len(output.tables)}, paragraphs={len(output.paragraphs)}")
print("PASS: first three table XML hashes are unchanged")
print("PASS: formulas and all required experiment/fitting values are present")
print("PASS: DOCX ZIP integrity is valid")
