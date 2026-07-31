from __future__ import annotations

import hashlib
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.shared import Pt


INPUT = Path(
    r"E:\DeskTop\preparation\Chen_Wei\TiaoZhanBei\Fitting\outputs\documents\5.5_反演结果不确定性与可信度评价_修订版.docx"
)
OUTPUT = Path(
    r"E:\DeskTop\preparation\Chen_Wei\TiaoZhanBei\Fitting\outputs\documents\5.5_反演结果不确定性与可信度评价_理论扩充版.docx"
)


def sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def table_hash(table: Table) -> str:
    return hashlib.sha256(table._tbl.xml.encode("utf-8")).hexdigest()


def set_run_font(run, equation: bool = False):
    latin = "Cambria Math" if equation else "Times New Roman"
    east = "Cambria Math" if equation else "宋体"
    run.font.name = latin
    run.font.size = Pt(11)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), east)
    rfonts.set(qn("w:cs"), latin)


def insert_after(anchor: Paragraph, text: str, equation: bool = False) -> Paragraph:
    element = OxmlElement("w:p")
    anchor._p.addnext(element)
    p = Paragraph(element, anchor._parent)
    p.style = anchor.part.document.styles["Normal"]
    if equation:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_together = True
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Pt(22)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, equation=equation)
    return p


def find_paragraph(doc: Document, needle: str) -> Paragraph:
    for p in doc.paragraphs:
        if needle in p.text:
            return p
    raise ValueError(f"Paragraph anchor not found: {needle}")


def insert_series(doc: Document, anchor_text: str, items: list[tuple[str, bool]]):
    anchor = find_paragraph(doc, anchor_text)
    for text, equation in items:
        anchor = insert_after(anchor, text, equation=equation)


def remove_section_56_table(doc: Document) -> int:
    in_section = False
    removed = 0
    for child in list(doc.element.body.iterchildren()):
        if isinstance(child, CT_P):
            p = Paragraph(child, doc)
            if p.text.startswith("5.5.6 "):
                in_section = True
            elif p.text.startswith("5.5.7 "):
                in_section = False
        elif isinstance(child, CT_Tbl) and in_section:
            doc.element.body.remove(child)
            removed += 1
    return removed


def build():
    input_hash = sha256(INPUT)
    doc = Document(INPUT)
    if len(doc.tables) < 4:
        raise ValueError(f"Expected at least 4 tables, found {len(doc.tables)}")
    preserved_hashes = [table_hash(t) for t in doc.tables[:3]]
    start_paragraphs = len([p for p in doc.paragraphs if p.text.strip()])

    insert_series(
        doc,
        "最终形成“多模型预测聚合",
        [
            (
                "从统计意义上看，不确定性并不等同于已经发生的预测误差。误差只有在真实标签已知时才能直接计算，而不确定性是在真实值尚不可得时，对潜在误差大小及排序关系的事前估计。因此，本节不以“给出一个看似精确的概率”为目标，而以能否提前识别高误差测点、能否为后续拟合提供合理的差异化权重为核心判据。",
                False,
            ),
            (
                "这一技术路线将模型预测、误差风险和工程权重划分为三个层次：位置级集成预测用于获得稳健状态基准，不确定性算法用于估计该基准的误差风险，经验分布映射再把风险排序转化为可供拟合和空间重构调用的权重。分层设计既保留了算法评价的统计含义，也避免下游模块直接依赖某一模型内部、难以统一解释的置信度尺度。",
                False,
            ),
        ],
    )

    insert_series(
        doc,
        "通常表现为原始响应的随机波动",
        [
            (
                "数据不确定性具有明显的条件相关性：其幅度会随钻进速度、接触状态、岩性、孔深和设备工作状态而变化，并非所有测点服从同一噪声分布。部分随机噪声即使增加样本数量也难以完全消除，属于观测过程中的固有波动；脉冲毛刺、短时失真和缺失值则可通过稳健检测与局部修正降低影响。因而，后续处理必须同时区分“可修正的异常”和“不可简单抹除的真实波动”。",
                False,
            ),
            (
                "沿孔深序列还具有较强的局部相关性。相邻窗口通常对应空间上连续的围岩区段，孤立单点突变更可能来自噪声，而连续多个测点维持的新水平则可能对应真实煤岩界面、破碎带或应力状态变化。该差异决定了本节在不确定性估计和拟合阶段均采用连续段保护，而不是把所有大幅变化统一视为坏数据。",
                False,
            ),
        ],
    )

    insert_series(
        doc,
        "模型间分歧不能单独作为可靠的不确定性指标",
        [
            (
                "模型不确定性又可理解为模型对训练样本覆盖范围和参数结构认识不足形成的认知不确定性。理论上，增加具有不同归纳偏置的模型、扩大训练工况覆盖并改善标签质量，可以降低该部分不确定性。多模型集成正是利用不同模型对同一输入的响应差异，近似观察模型空间中的预测波动；但这一近似成立的前提是模型之间具有足够多样性。",
                False,
            ),
            (
                "如果多个模型共享高度相似的数据、特征或决策边界，它们可能在同一位置同时给出相同但有偏的结果。此时“模型一致”只能说明模型输出相近，不能证明预测正确。因此，本研究除标准差和极差外，还引入模型平均置信度、累计孔深和区段编号，并直接利用历史绝对误差进行监督校准，以降低单一模型分歧指标失效带来的系统性风险。",
                False,
            ),
            (
                "在概率分解意义下，可将预测不确定性理解为数据条件波动与模型认知差异的共同作用。虽然当前数据不足以对二者进行严格概率分解，但这一理论框架说明：标准差、置信度、孔深和区段分别刻画了不同侧面的风险信息，只有在样本外实验中验证其与真实误差的对应关系后，才能将其用于工程赋权。",
                False,
            ),
            (
                "Var(y|x) = Eθ[Var(y|x,θ)] + Varθ[E(y|x,θ)]",
                True,
            ),
        ],
    )

    insert_series(
        doc,
        "作为不确定性算法的评价对象",
        [
            (
                "采用中位数而不是简单均值作为位置级参考预测，是因为当前每个测点仅有三个模型结果，单个模型的极端偏差可能明显拉动均值。中位数具有较高的抗异常能力，能够在至少两个模型结果相近时保持稳定，并为三个候选方案提供统一的误差基准。与此同时，标准差和极差仍被保留，用于描述集成内部的离散程度。",
                False,
            ),
            (
                "绝对误差作为监督目标具有直接的工程意义，其量纲与损伤或应力预测一致，不会因正负误差相互抵消。选择 90% 条件分位而非平均误差，意味着校准器更关注相似条件下误差分布的上部风险，能够减少少数大误差测点被平均水平掩盖的情况。这与后续安全约束中宁可对高风险位置增加平滑和降权、也不盲目保留异常预测的原则一致。",
                False,
            ),
        ],
    )

    insert_series(
        doc,
        "不解释为预测正确概率",
        [
            (
                "分位损失对低估和高估采用不对称惩罚。当目标分位数取 0.9 时，低估较大残差的代价高于高估同等幅度残差的代价，从而推动模型给出偏向误差上部的风险估计。与最小二乘回归预测条件均值相比，这种处理对偏斜、重尾和少量大误差更敏感，更适合复杂地质条件下的风险排序。",
                False,
            ),
            (
                "五项特征承担不同作用：标准差和极差描述模型间分歧，平均状态置信度描述模型内部的自评信息，累计孔深表征随钻过程和空间位置变化，区段编号则吸收不同数据段之间的系统差异。梯度提升树通过分段划分学习这些变量的非线性关系及交互作用，因此即使部分分歧特征退化，仍可能从孔深、区段和置信度中提取有效风险信号。",
                False,
            ),
            (
                "该方案的代价是输出具有树模型典型的分段常数特征。当大量样本落入相同叶节点时，不确定性会集中为少量档位。因而，方案一是否有效不能只看输出唯一值数量，而应同时考察 AURC、Spearman、不同数据源上的稳定性以及映射为权重后的动态范围。",
                False,
            ),
        ],
    )

    insert_series(
        doc,
        "不能直接利用孔深、区段及模型分歧等补充信息",
        [
            (
                "保序回归的核心作用是把工程先验写入校准函数：当模型自报的置信度降低时，估计误差风险不应反向下降。分箱后统计 90% 分位误差，可以减弱单点噪声对校准关系的影响；随后通过单调约束合并违背顺序的相邻区间，得到更平滑、可解释的风险函数。",
                False,
            ),
            (
                "这一方案同时是对 state_confidence 有效性的直接检验。如果单一置信度已经能够稳定排序误差，则复杂的多特征模型未必必要；反之，若置信度高度集中、缺失或与误差关系不单调，方案二的性能就会下降。当前缺失值按 0.5 的中性水平处理，能够避免样本被删除，但也会把来源不同的缺失原因压缩到同一风险位置，降低区分能力。",
                False,
            ),
        ],
    )

    insert_series(
        doc,
        "通过多源证据互补识别单一指标遗漏的风险位置",
        [
            (
                "融合前将三类证据转换为经验百分位，是为了消除标准差、置信缺失和局部偏离在数值尺度及量纲上的差异。百分位只保留各分量的相对排序，使融合系数可以解释为三类证据在综合风险中的相对贡献，而不会因某一分量数值范围较大而天然占优。",
                False,
            ),
            (
                "局部不稳定分量采用中位数和 MAD 而不是均值和标准差，是因为前者对孤立异常更稳健。但局部偏离本身并不能区分噪声与真实地质界面，因此又引入连续段和跨模型一致性保护。若变化持续不少于三个测点且多个模型在相近孔深同时响应，则降低该分量的异常贡献，以保留可能具有工程意义的边界和峰值。",
                False,
            ),
        ],
    )

    insert_series(
        doc,
        "不能把 273 条“系数组合×交叉验证折”记录直接视为 273 个独立折",
        [
            (
                "选择按 source_file 留一而不是随机拆分测点，是因为同一数据源中的相邻深度窗口往往具有相似工况、相同模型偏差和较强序列相关性。若随机划分，训练集和测试集可能同时包含同一钻孔或相邻窗口，评价结果容易因信息近邻而偏高。留一数据源验证要求算法迁移到完整未见数据源，更接近实际应用中的跨工况泛化。",
                False,
            ),
            (
                "AURC对应选择性预测思想：系统优先保留低不确定性样本，并逐步纳入高风险样本。若不确定性排序有效，低覆盖率阶段的平均误差应明显低于全体样本，风险曲线整体下移。AURC因此同时利用了多个覆盖水平的信息，比只在单一阈值上计算高低可信准确率更加稳定。",
                False,
            ),
            (
                "Spearman只评价排序的一致程度，对不确定性的绝对尺度和非线性单调变换不敏感。它能够补充说明“误差越大时风险是否总体越高”，但不能单独证明权重已经完成概率校准。AURC与Spearman共同使用，可以避免某一指标因误差量纲、并列值或少量极端样本而给出片面结论。",
                False,
            ),
        ],
    )

    insert_series(
        doc,
        "最终损伤和应力均选择条件分位残差回归",
        [
            (
                "从结果幅度看，方案一在损伤上的平均 AURC 比方案二降低约 31.1%，比方案三最佳固定组合降低约 18.1%；在应力上的平均 AURC 分别降低约 43.0% 和 43.0%。同时，方案一的损伤和应力 Spearman 均约为 0.38，明显高于其余两种方法，说明其优势并非只来自少量极端误差对 AURC 的影响。",
                False,
            ),
            (
                "方案二的落后说明原始状态置信度虽包含风险信息，但尚不足以独立承担误差校准任务。置信度高值集中使大量样本位于相近的风险区间，33.2% 的缺失又进一步压缩了可用信息。方案一将置信度作为多特征之一，可以在置信度失效或缺失时借助孔深、区段和集成统计进行补偿。",
                False,
            ),
            (
                "方案三的系数选择具有更直接的诊断意义。两个目标的最佳组合均令 γ=0，表明局部不稳定度与真实绝对误差之间缺少稳定的新增排序信息；应力最佳组合退化为纯置信缺失，也说明严重退化的模型分歧无法改善结果。多证据方案未胜出并非算法形式必然无效，而是实验证明当前三个证据源之间存在信息退化或重复。",
                False,
            ),
            (
                "需要指出，约 0.38 的 Spearman 属于中等相关水平，说明方案一能够提供有用但并非完美的风险排序。项目书中的结论应限定为“在当前数据和三套候选方法中表现最优”，而不是宣称已经准确刻画全部预测不确定性。该限定也为后续增加模型多样性和完善应力校准保留了清晰空间。",
                False,
            ),
        ],
    )

    insert_series(
        doc,
        "更适合解释为相对风险分级",
        [
            (
                "经验分布映射的主要作用是统一损伤和应力不确定性的尺度。原始不确定性保持目标量纲，而百分位将其转换为 0～1 范围内的相对位置，使同一套拟合公式可以处理两类目标。该变换保持风险排序，却不会自动使置信度成为统计意义上的覆盖概率或正确率。",
                False,
            ),
            (
                "并列值会直接影响权重分辨率。当前采用右侧经验累积分布，同一不确定性档位的测点被赋予相同百分位；当最低不确定性档位包含大量测点时，其映射权重也可能明显低于 1。应力最高权重仅为 0.567 正是这一机制与七档离散输出共同作用的结果，而不是模型分歧大的证据。",
                False,
            ),
            (
                "分别构造 damage_weight 和 stress_weight 具有必要性。损伤反映结构劣化和破坏程度，应力反映承载水平和集中状态，两者的误差分布、可辨识性和校准分辨率均不相同。后续拟合及 5.6 空间重构应分别使用对应权重，只有在风险判识阶段才根据工程目标综合解释两类状态场。",
                False,
            ),
        ],
    )

    insert_series(
        doc,
        "Iᵢ = 1{|yᵢ−mᵢ|",
        [
            (
                "Hampel 判别以局部中位数作为中心，以 MAD 作为稳健尺度。当窗口内存在少量极端值时，中位数和 MAD 不会像均值与标准差那样被异常值显著拉动，因此更适合沿孔深序列的孤立毛刺识别。阈值 3.5×1.4826MAD 将 MAD 近似换算到正态标准差尺度，但连续段保护进一步避免对真实结构变化进行机械修正。",
                False,
            ),
        ],
    )

    insert_series(
        doc,
        "ŷᵢ,fit⁽ᵏ⁾ =",
        [
            (
                "Savitzky-Golay 方法在滑动窗口内使用低阶多项式进行局部最小二乘拟合。与简单移动平均相比，它在抑制高频波动的同时能够更好地保持峰值位置、斜率和局部曲率。权重融合不是直接改变滤波窗口，而是在修正后原曲线与统一平滑参考之间连续插值，使不确定性方案与平滑算法保持解耦。",
                False,
            ),
            (
                "由 λᵢ⁽ᵏ⁾=0.8−0.6wᵢ⁽ᵏ⁾ 可知，当权重为 0.1 时，平滑融合比例为 0.74；当权重达到 1.0 时，平滑融合比例降为 0.20。权重越高，结果越接近修正后原值；权重越低，结果越依赖平滑曲线。应力当前最高权重为 0.567，对应最小平滑融合比例约为 0.46，这也解释了应力曲线整体仍保持较强平滑。",
                False,
            ),
        ],
    )

    insert_series(
        doc,
        "形成用于后续空间重构的连续损伤曲线和连续应力曲线",
        [
            (
                "PCHIP 在相邻测点之间构造分段三次 Hermite 多项式，并利用相邻割线斜率限制节点导数。当原序列在局部保持单调时，PCHIP 通常不会像普通三次样条那样产生新的极值和明显过冲，适合损伤边界和应力峰值的致密表达。拟合后还对损伤、应力施加原组数据范围约束，进一步保证输出满足物理取值边界。",
                False,
            ),
            (
                "曲线处理效果通过总变差和粗糙度从不同角度评价。总变差统计全部相邻变化的绝对值，较小通常表示曲线整体波动减弱；粗糙度统计相邻一阶差分变化，反映高频振荡和局部折线程度。二者均不能单独越小越好，因为过度平滑也会压低真实峰值和边界，必须结合峰值保持、范围过冲及真实误差共同判断。",
                False,
            ),
            ("TV(y) = Σᵢ₌₁ᴺ⁻¹ |yᵢ₊₁ − yᵢ|", True),
            ("Roughness(y) = meanᵢ{|(yᵢ₊₂−yᵢ₊₁) − (yᵢ₊₁−yᵢ)|}", True),
        ],
    )

    insert_series(
        doc,
        "尚不能据此宣称加权方案在全量数据上取得显著提升",
        [
            (
                "从 roughness 看，损伤有权重和无权重结果分别为 0.554 和 0.204，应力分别为 0.191 和 0.104。无权重结果更平滑，但这并不自动意味着更优，因为其损伤总变差达到 466，偏离原始 TV=420 的程度更大。有权重模式在损伤上保留了更多局部变化，因此表现为粗糙度略高而总变差更接近原始序列，体现了保真与平滑之间的权衡。",
                False,
            ),
            (
                "应力原始 TV 为 680，两种模式均将其降低至约 4.1×10²，说明该组应力序列需要较强的高频抑制；有权重的 419 与无权重的 413 相差较小。所有结果均未出现超出原始取值范围的过冲，表明范围约束和 PCHIP 保形插值能够控制拟合稳定性。由于目前只有单组对比，评价重点仍是流程可运行、权重作用方向符合设计，而非统计显著性。",
                False,
            ),
        ],
    )

    insert_series(
        doc,
        "现有数据的信号边界",
        [
            (
                "从科研设计角度看，方案二和方案三构成了必要的消融对照。方案二回答“模型自带置信度是否足够”，方案三回答“增加模型分歧和局部稳定性是否能够进一步改善”。两者均未胜出，使方案一的选择不再只是经验偏好，而是建立在统一数据划分、统一指标和统一目标下的比较结果。",
                False,
            ),
        ],
    )

    insert_series(
        doc,
        "后续仍需在全部模型—数据源—实验分组上开展配对统计",
        [
            (
                "模型多样性是后续改进的首要方向。可通过引入不同特征体系、不同模型结构或不同训练样本重采样，增加集成成员之间的有效差异，使模型间标准差能够真实反映认知边界，而不是大量退化为零。新增模型后仍应采用同样的留一数据源验证，确认分歧增加对应的是误差信息而非无效随机波动。",
                False,
            ),
            (
                "应力权重的改进重点是提高条件风险的分辨率并妥善处理并列值。可比较更高容量的分位模型、以训练折为参照的中位秩映射和分段校准方法，但任何拉伸权重范围的处理都必须先验证 AURC 和样本外误差，不能只为了获得接近 1 的权重而人为放大差异。",
                False,
            ),
            (
                "全量拟合验证应在相同分组和相同拟合参数下，对有权重、无权重结果进行配对比较，至少同时报告相对真实标签的 MAE/RMSE、总变差、最大跳变、粗糙度、峰值保持率和过冲。只有当多数分组获得稳定改善且配对差异具有统计支持时，才能进一步表述为不确定性权重提高了整体拟合性能。",
                False,
            ),
        ],
    )

    insert_series(
        doc,
        "避免将单孔模型风险与空间外推风险混为同一指标",
        [
            (
                "在 5.6 的可信度—距离联合权重中，本节输出的 damage_weight 和 stress_weight 只代表观测值本身的反演可靠程度。二维或三维待估位置还需结合观测距离、邻域独立钻孔数量、周向覆盖均衡性、径向孔深覆盖、轨迹定位精度及插值模型方差重新评价。两层可信度分开构建，可避免把“测点预测可靠”误解为“远距离空间外推同样可靠”。",
                False,
            ),
            (
                "因此，5.5 的最终产物不是单一工程等级，而是一组能够随目标和孔深变化的连续权重；5.6 则在这些权重基础上形成空间可信度场。该接口既允许高可信测点在邻近区域发挥较大作用，也允许空间覆盖不足区域即使邻近测点本身可靠，仍因外推距离和几何约束而降低空间可信度。",
                False,
            ),
        ],
    )

    removed = remove_section_56_table(doc)
    if removed != 1:
        raise ValueError(f"Expected to remove exactly one table in 5.5.6, removed {removed}")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = "5.5 反演结果不确定性与可信度评价（理论扩充版）"
    doc.save(OUTPUT)

    if sha256(INPUT) != input_hash:
        raise AssertionError("Input document changed")
    revised = Document(OUTPUT)
    final_hashes = [table_hash(t) for t in revised.tables]
    if final_hashes != preserved_hashes:
        raise AssertionError(f"Preserved table hashes changed: {preserved_hashes} != {final_hashes}")
    final_paragraphs = len([p for p in revised.paragraphs if p.text.strip()])
    print(f"Created: {OUTPUT}")
    print(f"Input SHA256 preserved: {input_hash}")
    print(f"Tables preserved: {len(final_hashes)}")
    print(f"Paragraphs: {start_paragraphs} -> {final_paragraphs}")
    print(f"Output bytes: {OUTPUT.stat().st_size}")


if __name__ == "__main__":
    build()
